#!/usr/bin/env python3
"""Minimal Markdown to .docx converter for this project's articles.

Handles the subset of Markdown the articles use: YAML frontmatter (stripped),
ATX headings, blockquotes, paragraphs, and inline [text](url) links rendered as
real Word hyperlinks. HTML comment blocks are dropped.
"""
import os
import re
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs_with_inline(paragraph, text):
    """Render text with inline links and bold into a paragraph."""
    pos = 0
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            _add_plain_or_bold(paragraph, text[pos:m.start()])
        add_hyperlink(paragraph, m.group(2), m.group(1))
        pos = m.end()
    if pos < len(text):
        _add_plain_or_bold(paragraph, text[pos:])


def _add_plain_or_bold(paragraph, text):
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def strip_frontmatter(lines):
    if lines and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                return lines[i + 1:]
    return lines


def main(src, dst):
    with open(src, encoding="utf-8") as f:
        lines = f.read().splitlines()
    lines = strip_frontmatter(lines)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    NAVY = RGBColor(0x1F, 0x3A, 0x5F)

    def heading(text, level, size):
        h = doc.add_heading(level=level)
        run = h.add_run(text)
        run.font.color.rgb = NAVY
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        return h

    in_comment = False
    buf = []
    mode = None  # "para" or "quote"

    def flush():
        nonlocal buf, mode
        if not buf:
            return
        text = " ".join(buf)
        if mode == "quote":
            p = doc.add_paragraph(style="Intense Quote")
        else:
            p = doc.add_paragraph()
        add_runs_with_inline(p, text)
        buf = []
        mode = None

    for line in lines:
        stripped = line.strip()
        if "<!--" in stripped:
            flush()
            in_comment = True
        if in_comment:
            if "-->" in stripped:
                in_comment = False
            continue
        if not stripped or stripped == "---":
            flush()
            continue
        m_img = IMAGE_RE.match(stripped)
        if m_img:
            flush()
            alt, path = m_img.group(1), m_img.group(2)
            img_path = path if os.path.isabs(path) else os.path.join(
                os.path.dirname(os.path.abspath(src)), path
            )
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Inches(6.2))
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.add_run(alt)
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
            continue
        if stripped.startswith("### "):
            flush()
            heading(stripped[4:].strip(), 2, 13)
        elif stripped.startswith("## "):
            flush()
            heading(stripped[3:].strip(), 1, 15)
        elif stripped.startswith("# "):
            flush()
            heading(stripped[2:].strip(), 0, 26)
        elif stripped.startswith(">"):
            if mode != "quote":
                flush()
                mode = "quote"
            buf.append(stripped.lstrip("> ").strip())
        else:
            if mode != "para":
                flush()
                mode = "para"
            buf.append(stripped)

    flush()
    doc.save(dst)
    print("wrote", dst)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
